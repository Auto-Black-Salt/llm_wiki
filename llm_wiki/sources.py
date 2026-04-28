import hashlib
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import httpx
from readability import Document
from lxml import html as lxml_html

MAX_CHARS = 50_000  # ~12k tokens, safe for most context windows


@dataclass
class ParsedSource:
    filename: str
    text: str
    raw_bytes: Optional[bytes] = None
    images: list[tuple[str, bytes]] = field(default_factory=list)  # (filename, data)


def parse_source(path_or_url: str) -> ParsedSource:
    """Parse a file path or URL into a ParsedSource."""
    if path_or_url.startswith(("http://", "https://")):
        if _is_youtube_url(path_or_url):
            return _fetch_youtube(path_or_url)
        return _fetch_url(path_or_url)
    path = Path(path_or_url)
    if path.suffix.lower() == ".pdf":
        return _parse_pdf(path)
    if path.suffix.lower() in (".docx", ".doc"):
        return _parse_docx(path)
    return ParsedSource(filename=path.name, text=path.read_text())


def chunk_text(text: str, max_chars: int = MAX_CHARS) -> list[str]:
    """Split text into chunks of at most max_chars, breaking at paragraph boundaries."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    while text:
        if len(text) <= max_chars:
            chunks.append(text)
            break
        split_at = text.rfind("\n\n", 0, max_chars)
        if split_at == -1:
            split_at = text.rfind("\n", 0, max_chars)
        if split_at == -1:
            split_at = max_chars
        chunks.append(text[:split_at])
        text = text[split_at:].lstrip("\n")
    return chunks


def _is_youtube_url(url: str) -> bool:
    return "youtube.com/watch" in url or "youtu.be/" in url


def _extract_video_id(url: str) -> str:
    if "youtu.be/" in url:
        return url.split("youtu.be/")[1].split("?")[0].split("/")[0]
    from urllib.parse import urlparse, parse_qs
    return parse_qs(urlparse(url).query)["v"][0]


def _fetch_youtube(url: str) -> ParsedSource:
    from youtube_transcript_api import YouTubeTranscriptApi  # lazy import

    video_id = _extract_video_id(url)

    # Fetch video title from page metadata
    try:
        resp = httpx.get(url, follow_redirects=True, timeout=15)
        title_match = re.search(r'<meta property="og:title" content="([^"]+)"', resp.text)
        title = title_match.group(1) if title_match else video_id
    except Exception:
        title = video_id

    # Fetch transcript — try all available languages if English not found
    api = YouTubeTranscriptApi()
    try:
        transcript = api.fetch(video_id)
    except Exception:
        try:
            transcript_list = api.list(video_id)
            transcript = transcript_list.find_transcript(
                [t.language_code for t in transcript_list]
            ).fetch()
        except Exception as e:
            raise ValueError(f"No transcript available for {url}: {e}") from e

    text = "\n".join(snippet.text for snippet in transcript)
    slug = re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "-")[:60]
    filename = f"{slug}-{video_id}.md"
    full_text = f"# {title}\n\n**YouTube:** {url}\n\n## Transcript\n\n{text}"
    return ParsedSource(filename=filename, text=full_text)


def _fetch_url(url: str) -> ParsedSource:
    response = httpx.get(url, follow_redirects=True, timeout=30)
    response.raise_for_status()
    doc = Document(response.text)
    summary_html = doc.summary()
    tree = lxml_html.fromstring(summary_html)
    text = tree.text_content()
    if not text.strip():
        raise ValueError(f"Could not extract readable content from {url}")
    slug = re.sub(r"[^\w-]", "-", url.split("//")[-1].split("/")[0])[:40]
    hash_suffix = hashlib.md5(url.encode()).hexdigest()[:4]
    filename = f"{slug}-{hash_suffix}.html"
    return ParsedSource(filename=filename, text=text, raw_bytes=response.content)


def _parse_docx(path: Path) -> ParsedSource:
    import docx  # lazy import — only needed for Word docs
    doc = docx.Document(str(path))
    parts = []
    images = []
    img_counter = 0

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Extract images from the document's inline shapes / relationships
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            ext = rel.target_part.content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            img_counter += 1
            img_filename = f"{path.stem}-img{img_counter}.{ext}"
            images.append((img_filename, img_data))
            parts.append(f"![[assets/{img_filename}]]")

    return ParsedSource(filename=path.name, text="\n".join(parts), images=images)


def _parse_pdf(path: Path) -> ParsedSource:
    try:
        import opendataloader_pdf  # lazy import — only needed for PDFs
    except ModuleNotFoundError as e:
        raise ValueError(
            "PDF parsing requires the 'opendataloader-pdf' package. "
            "Install it and make sure Java 11+ is available."
        ) from e

    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            opendataloader_pdf.convert(
                input_path=[str(path)],
                output_dir=tmpdir,
                format="markdown",
                image_output="off",
                use_struct_tree=True,
                quiet=True,
            )
        except FileNotFoundError as e:
            raise ValueError(
                "OpenDataLoader PDF requires Java 11+. "
                "Install a JDK and ensure 'java' is on PATH."
            ) from e
        except Exception as e:
            raise ValueError(f"Could not convert PDF {path} with OpenDataLoader PDF: {e}") from e

        output_dir = Path(tmpdir)
        markdown_path = _find_markdown_output(output_dir, path.stem)
        if markdown_path is None:
            raise ValueError(f"OpenDataLoader PDF did not produce Markdown output for {path}")

        return ParsedSource(filename=path.name, text=markdown_path.read_text())


def _find_markdown_output(output_dir: Path, stem: str) -> Optional[Path]:
    """Find the Markdown file produced for a converted PDF."""
    exact_matches = sorted(output_dir.rglob(f"{stem}.md"))
    if exact_matches:
        return exact_matches[0]

    markdown_files = sorted(output_dir.rglob("*.md"))
    if len(markdown_files) == 1:
        return markdown_files[0]
    if markdown_files:
        return markdown_files[0]
    return None
